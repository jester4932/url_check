import re

from docx import Document #python-docx


class DOCXRead:
    def __init__(self, **kwargs):
        self.file = kwargs.get("file", "")
        self.regex = kwargs.get("regex", "")
        self.doc_parts = [
            "webSettings.xml",
            "settings.xml",
            "styles.xml",
            "theme/theme1.xml",
            "fontTable.xml",
        ]
        self.doc = Document(self.file)
        self.url = []

    def main(self):
        self.plain_text()
        self.hyperlinks()
        return self.url

    def plain_text(self):  # find plain text url
        all_text = []
        for para in self.doc.paragraphs:
            all_text.append(para.text)
        for text in all_text:
            url = re.findall(self.regex, text)
            if url:
                self.url.append(text)

    def hyperlinks(self):  # find hyperlinks
        first_line = ""
        url_pattern = re.compile(self.regex)
        if self.doc.paragraphs:
            first_line = self.doc.paragraphs[0]
        for rel in first_line.part.rels.values():
            # look for embedded hyperlink on top line of doc
            target = rel.target_ref
            if (
                url_pattern.match(target)
                and not target.startswith("mailto:")
                and target not in self.doc_parts
            ):
                self.url.append(target)
        for rel in self.doc.part.rels.values():
            target = rel.target_ref
            if (
                url_pattern.match(target)
                and not target.startswith("mailto:")
                and target not in self.doc_parts
            ):
                self.url.append(target)
